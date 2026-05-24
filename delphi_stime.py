#!/usr/bin/env python3
"""
Raccolta stime Delphi per il progetto EvenToNight.
Stime in giorni/uomo. Multi-round per attività, 3 partecipanti.
Salva lo stato in JSON e genera un Markdown strutturato.
"""

import json
import math
import os
import statistics
from datetime import datetime, timedelta

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── ATTIVITÀ FOGLIA della WBS ──────────────────────────────────────────────
# (id, nome, percorso_padre)

ACTIVITIES = [
    # ── 1. MARKETING ──────────────────────────────────────────────────────
    ("1.1.1.1",  "Ricerca Organizzazioni",
     "1 Marketing > 1.1 Coinvolgimento Organizzazioni > 1.1.1 Gestione Acquisizione Nuove Organizzazioni"),
    ("1.1.1.2",  "Contatto Organizzazioni",
     "1 Marketing > 1.1 Coinvolgimento Organizzazioni > 1.1.1 Gestione Acquisizione Nuove Organizzazioni"),
    ("1.1.1.3",  "Onboarding Organizzazioni",
     "1 Marketing > 1.1 Coinvolgimento Organizzazioni > 1.1.1 Gestione Acquisizione Nuove Organizzazioni"),
    ("1.1.2",    "Gestione Contatto Organizzazioni già Registrate",
     "1 Marketing > 1.1 Coinvolgimento Organizzazioni"),
    ("1.2.1.1",  "Individuazione Profilo medio utente",
     "1 Marketing > 1.2 Coinvolgimento Utenti > 1.2.1 Analisi Dati Utenti Attivi"),
    ("1.2.1.2",  "Analisi tipologia di Eventi più remunerativi",
     "1 Marketing > 1.2 Coinvolgimento Utenti > 1.2.1 Analisi Dati Utenti Attivi"),
    ("1.2.1.3",  "Analisi Provenienza geografica delle visualizzazioni della piattaforma",
     "1 Marketing > 1.2 Coinvolgimento Utenti > 1.2.1 Analisi Dati Utenti Attivi"),
    ("1.2.2.1.1","Ricerca e selezione influencer",
     "1 Marketing > 1.2 Coinvolgimento Utenti > 1.2.2 Campagna Pubblicitaria > 1.2.2.1 Influencer Marketing"),
    ("1.2.2.1.2","Contatto e accordi",
     "1 Marketing > 1.2 Coinvolgimento Utenti > 1.2.2 Campagna Pubblicitaria > 1.2.2.1 Influencer Marketing"),
    ("1.2.2.1.3","Briefing contenuti",
     "1 Marketing > 1.2 Coinvolgimento Utenti > 1.2.2 Campagna Pubblicitaria > 1.2.2.1 Influencer Marketing"),
    ("1.2.2.1.4","Monitoraggio e ottimizzazione pubblicazioni",
     "1 Marketing > 1.2 Coinvolgimento Utenti > 1.2.2 Campagna Pubblicitaria > 1.2.2.1 Influencer Marketing"),
    ("1.2.2.2.1","Ricerca e selezione piattaforme ads",
     "1 Marketing > 1.2 Coinvolgimento Utenti > 1.2.2 Campagna Pubblicitaria > 1.2.2.2 ADS online"),
    ("1.2.2.2.2","Creazione contenuti (grafiche/video)",
     "1 Marketing > 1.2 Coinvolgimento Utenti > 1.2.2 Campagna Pubblicitaria > 1.2.2.2 ADS online"),
    ("1.2.2.2.3","Setup campagna",
     "1 Marketing > 1.2 Coinvolgimento Utenti > 1.2.2 Campagna Pubblicitaria > 1.2.2.2 ADS online"),
    ("1.2.2.2.4","Monitoraggio e ottimizzazione",
     "1 Marketing > 1.2 Coinvolgimento Utenti > 1.2.2 Campagna Pubblicitaria > 1.2.2.2 ADS online"),

    # ── 2. APP MOBILE ─────────────────────────────────────────────────────
    # 2.1.1 Home
    ("2.1.1.1.1","Design della schermata Home per utenti non registrati",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.1 Home"),
    ("2.1.1.1.2","Implementazione della schermata Home per utenti non registrati",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.1 Home"),
    ("2.1.1.2.1","Design della schermata Home per utenti registrati",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.1 Home"),
    ("2.1.1.2.2","Implementazione della schermata Home per utenti registrati",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.1 Home"),
    # 2.1.2 Esplora
    ("2.1.2.1.1","Design del sistema di ricerca per la schermata Esplora",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.2 Esplora"),
    ("2.1.2.1.2","Implementazione del sistema di ricerca per la schermata Esplora",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.2 Esplora"),
    ("2.1.2.2.1","Design della schermata Esplora per utenti non registrati",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.2 Esplora"),
    ("2.1.2.2.2","Implementazione della schermata Esplora per utenti non registrati",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.2 Esplora"),
    ("2.1.2.3.1","Design della schermata Esplora per utenti registrati",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.2 Esplora"),
    ("2.1.2.3.2","Implementazione della schermata Esplora per utenti registrati",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.2 Esplora"),
    ("2.1.2.4.1","Design dei filtri per gli eventi nella schermata Esplora",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.2 Esplora"),
    ("2.1.2.4.2","Implementazione dei filtri per gli eventi nella schermata Esplora",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.2 Esplora"),
    # 2.1.3 Pagina dell'evento
    ("2.1.3.1.1","Design della schermata Pagina Evento",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.3 Pagina dell'evento"),
    ("2.1.3.1.2","Implementazione della schermata Pagina Evento",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.3 Pagina dell'evento"),
    ("2.1.3.2.1","Design della sezione Recensioni",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.3 Pagina dell'evento"),
    ("2.1.3.2.2","Implementazione della sezione Recensioni",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.3 Pagina dell'evento"),
    ("2.1.3.3.1","Design della navigazione verso la schermata di checkout",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.3 Pagina dell'evento"),
    ("2.1.3.3.2","Implementazione della schermata di checkout",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.3 Pagina dell'evento"),
    ("2.1.3.4.1","Design della funzionalità di like",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.3 Pagina dell'evento"),
    ("2.1.3.4.2","Implementazione della funzionalità di like",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.3 Pagina dell'evento"),
    # 2.1.4 Checkout
    ("2.1.4.1.1","Design della sezione dei biglietti nella schermata di checkout",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.4 Pagina di checkout"),
    ("2.1.4.1.2","Implementazione della sezione dei biglietti nella schermata di checkout",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.4 Pagina di checkout"),
    ("2.1.4.3.1","Design del sistema di completamento dell'acquisto",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.4 Pagina di checkout"),
    ("2.1.4.3.2","Implementazione del sistema di completamento dell'acquisto",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.4 Pagina di checkout"),
    # 2.1.5 Recensioni
    ("2.1.5.1.1","Design della schermata delle recensioni",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.5 Recensioni"),
    ("2.1.5.1.2","Implementazione della schermata delle recensioni",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.5 Recensioni"),
    ("2.1.5.2.1","Design del resoconto del rating dell'organizzazione",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.5 Recensioni"),
    ("2.1.5.2.2","Implementazione del resoconto del rating",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.5 Recensioni"),
    ("2.1.5.3.1","Design del sistema di inserimento di una recensione",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.5 Recensioni"),
    ("2.1.5.3.2","Implementazione del sistema di inserimento di una recensione",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.5 Recensioni"),
    # 2.1.6 Profilo
    ("2.1.6.1.1", "Design della sezione delle informazioni del profilo",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.1.2", "Implementazione della sezione delle informazioni del profilo",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.2.1", "Design del sistema di modifica del profilo",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.2.2", "Implementazione del sistema di modifica del profilo",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.3.1", "Design del sistema di modifica delle impostazioni dell'account",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.3.2", "Implementazione del sistema di modifica delle impostazioni dell'account",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.4.1", "Design del sistema di follow/unfollow",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.4.2", "Implementazione del sistema di follow/unfollow",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.5.1", "Design del sistema di contatto nei profili delle organizzazioni",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.5.2", "Implementazione del sistema di contatto nei profili delle organizzazioni",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.6.1", "Design della sezione degli eventi a cui l'utente ha messo like",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.6.2", "Implementazione della sezione degli eventi a cui l'utente ha messo like",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.7.1", "Design della sezione degli eventi a cui l'utente ha partecipato",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.7.2", "Implementazione della sezione degli eventi a cui l'utente ha partecipato",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.8.1", "Design della sezione degli eventi creati da un'organizzazione",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.8.2", "Implementazione della sezione degli eventi creati da un'organizzazione",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.9.1", "Design della sezione delle bozze degli eventi",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.9.2", "Implementazione della sezione delle bozze degli eventi",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.10.1","Design del sistema di visualizzazione delle statistiche",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.10.2","Implementazione del sistema di visualizzazione delle statistiche",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.11.1","Design della sezione delle recensioni",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    ("2.1.6.11.2","Implementazione della sezione delle recensioni",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.6 Profilo"),
    # 2.1.7 Biglietti
    ("2.1.7.1.1","Design della schermata dei biglietti",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.7 Biglietti"),
    ("2.1.7.1.2","Implementazione della schermata dei biglietti",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.7 Biglietti"),
    # 2.1.8 Chat
    ("2.1.8.1.1","Design della sezione delle conversazioni",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.8 Chat"),
    ("2.1.8.1.2","Implementazione della sezione delle conversazioni",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.8 Chat"),
    ("2.1.8.2.1","Design della schermata delle chat",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.8 Chat"),
    ("2.1.8.2.2","Implementazione della schermata delle chat",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.8 Chat"),
    ("2.1.8.3.1","Design della funzionalità di ricerca delle conversazioni",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.8 Chat"),
    ("2.1.8.3.2","Implementazione della schermata di ricerca delle conversazioni",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.8 Chat"),
    # 2.1.9 Creazione Evento
    ("2.1.9.1.1","Design della sezione delle informazioni dell'evento",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.9 Creazione Evento"),
    ("2.1.9.1.2","Implementazione della sezione delle informazioni dell'evento",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.9 Creazione Evento"),
    ("2.1.9.2.1","Design delle funzionalità di gestione dell'evento",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.9 Creazione Evento"),
    ("2.1.9.2.2","Implementazione delle funzionalità di gestione dell'evento",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.9 Creazione Evento"),
    # 2.1.10 Registrazione/Login
    ("2.1.10.1.1","Design della schermata di registrazione/login",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.10 Registrazione/Login"),
    ("2.1.10.1.2","Implementazione della schermata di registrazione/login",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.10 Registrazione/Login"),
    # 2.1.11 Scansione biglietti
    ("2.1.11.1.1","Design della schermata di scansione biglietti",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.11 Scansione biglietti"),
    ("2.1.11.1.2","Implementazione della schermata di scansione biglietti",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.11 Scansione biglietti"),
    # 2.1.12 Navigazione
    ("2.1.12.1.1","Design del sistema di navigazione",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.12 Sistema di navigazione"),
    ("2.1.12.1.2","Implementazione del sistema di navigazione",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.12 Sistema di navigazione"),
    # 2.1.13 Requisiti Non Funzionali
    ("2.1.13.1.1","Verificare che l'app sia responsive",
     "2 App Mobile > 2.1 Interfaccia Utente > 2.1.13 Requisiti Non Funzionali"),
    # 2.2 Notifiche push
    ("2.2.1","Studio Funzionamento Notifiche push",
     "2 App Mobile > 2.2 Notifiche push"),
    ("2.2.2","Implementazione Notifiche push",
     "2 App Mobile > 2.2 Notifiche push"),
    # 2.3 Testing UX
    ("2.3.1.1","Individuazione dei Principali Task Per Valutare l'usabilità",
     "2 App Mobile > 2.3 Testing UX"),
    ("2.3.1.2","Contatto Utenti Target e Organizzazione Sessioni di Usability Test",
     "2 App Mobile > 2.3 Testing UX"),
    ("2.3.1.3","Svolgimento Sessioni di Usability Test",
     "2 App Mobile > 2.3 Testing UX"),
    ("2.3.1.4","Valutazione Risultati Dei Test Di Usabilità",
     "2 App Mobile > 2.3 Testing UX"),
    # 2.4 Pubblicazione
    ("2.4.1.1","Configurazione account Google Play",
     "2 App Mobile > 2.4 Pubblicazione"),
    ("2.4.1.2","Configurazione account App Store",
     "2 App Mobile > 2.4 Pubblicazione"),
    ("2.4.1.3","Preparazione asset (icone, screenshot, descrizioni)",
     "2 App Mobile > 2.4 Pubblicazione"),
    ("2.4.1.4","Build e firma dell'app",
     "2 App Mobile > 2.4 Pubblicazione"),
    ("2.4.1.5","Submission e gestione review",
     "2 App Mobile > 2.4 Pubblicazione"),

    # ── 3. AGGIORNAMENTO SISTEMA ESISTENTE ────────────────────────────────
    # 3.1 OAuth2
    ("3.1.1","Scelta provider OAuth2 (e.g. Google)",
     "3 Aggiornamento Sistema Esistente > 3.1 Registrazione con OAuth2"),
    ("3.1.2","Aggiornamento configurazione servizio Autenticazione",
     "3 Aggiornamento Sistema Esistente > 3.1 Registrazione con OAuth2"),
    ("3.1.3","Testing servizio Autenticazione",
     "3 Aggiornamento Sistema Esistente > 3.1 Registrazione con OAuth2"),
    ("3.1.4","Integrazione aggiornamenti nel servizio Utenti",
     "3 Aggiornamento Sistema Esistente > 3.1 Registrazione con OAuth2"),
    ("3.1.5","Aggiornamento Documentazione API del servizio Utenti",
     "3 Aggiornamento Sistema Esistente > 3.1 Registrazione con OAuth2"),
    ("3.1.6","Testing servizio Utenti",
     "3 Aggiornamento Sistema Esistente > 3.1 Registrazione con OAuth2"),
    # 3.2 Geolocalizzazione
    ("3.2.1.1","Sviluppo endpoint RBS 3.2.1 nel servizio Geolocalizzazione",
     "3 Aggiornamento Sistema Esistente > 3.2 Servizio Geolocalizzazione"),
    ("3.2.1.2","Documentazione API con endpoint RBS 3.2.1",
     "3 Aggiornamento Sistema Esistente > 3.2 Servizio Geolocalizzazione"),
    ("3.2.1.3","Testing della funzionalità RBS 3.2.1",
     "3 Aggiornamento Sistema Esistente > 3.2 Servizio Geolocalizzazione"),
    ("3.2.2.1","Sviluppo endpoint RBS 3.2.2 nel servizio Geolocalizzazione",
     "3 Aggiornamento Sistema Esistente > 3.2 Servizio Geolocalizzazione"),
    ("3.2.2.2","Integrazione campo posizizione in tutti i servizi che lo richiedono",
     "3 Aggiornamento Sistema Esistente > 3.2 Servizio Geolocalizzazione"),
    ("3.2.2.3","Documentazione API con endpoint RBS 3.2.2",
     "3 Aggiornamento Sistema Esistente > 3.2 Servizio Geolocalizzazione"),
    ("3.2.2.4","Testing della funzionalità RBS 3.2.2",
     "3 Aggiornamento Sistema Esistente > 3.2 Servizio Geolocalizzazione"),
    # 3.3 Notifiche (8 gruppi × 3 task)
    ("3.3.1.1","Aggiunta notifica RBS 3.3.1 al servizio Notifiche",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.1.2","Aggiornamento Documentazione API con notifica RBS 3.3.1",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.1.3","Testing della funzionalità di notifica RBS 3.3.1",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.2.1","Aggiunta notifica RBS 3.3.2 al servizio Notifiche",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.2.2","Aggiornamento Documentazione API con notifica RBS 3.3.2",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.2.3","Testing della funzionalità di notifica RBS 3.3.2",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.3.1","Aggiunta notifica RBS 3.3.3 al servizio Notifiche",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.3.2","Aggiornamento Documentazione API con notifica RBS 3.3.3",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.3.3","Testing della funzionalità di notifica RBS 3.3.3",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.4.1","Aggiunta notifica RBS 3.3.4 al servizio Notifiche",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.4.2","Aggiornamento Documentazione API con notifica RBS 3.3.4",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.4.3","Testing della funzionalità di notifica RBS 3.3.4",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.5.1","Aggiunta notifica RBS 3.3.5 al servizio Notifiche",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.5.2","Aggiornamento Documentazione API con notifica RBS 3.3.5",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.5.3","Testing della funzionalità di notifica RBS 3.3.5",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.6.1","Aggiunta notifica RBS 3.3.6 al servizio Notifiche",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.6.2","Aggiornamento Documentazione API con notifica RBS 3.3.6",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.6.3","Testing della funzionalità di notifica RBS 3.3.6",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.7.1","Aggiunta notifica RBS 3.3.7 al servizio Notifiche",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.7.2","Aggiornamento Documentazione API con notifica RBS 3.3.7",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.7.3","Testing della funzionalità di notifica RBS 3.3.7",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.8.1","Aggiunta notifica RBS 3.3.8 al servizio Notifiche",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.8.2","Aggiornamento Documentazione API con notifica RBS 3.3.8",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    ("3.3.8.3","Testing della funzionalità di notifica RBS 3.3.8",
     "3 Aggiornamento Sistema Esistente > 3.3 Servizio Notifiche"),
    # 3.4 Abbonamento
    ("3.4.1","Sviluppo servizio Subscription per la gestione dei piani di abbonamento",
     "3 Aggiornamento Sistema Esistente > 3.4 Piani di Abbonamento"),
    ("3.4.2","Integrazione servizio Subscription con il resto del sistema",
     "3 Aggiornamento Sistema Esistente > 3.4 Piani di Abbonamento"),
    ("3.4.3","Documentazione API del servizio Subscription",
     "3 Aggiornamento Sistema Esistente > 3.4 Piani di Abbonamento"),
    ("3.4.4","Testing servizio Subscription",
     "3 Aggiornamento Sistema Esistente > 3.4 Piani di Abbonamento"),
    # 3.5 Ticketing
    ("3.5.1.1","Aggiunta endpoint RBS 3.5.1 al servizio Ticketing",
     "3 Aggiornamento Sistema Esistente > 3.5 Servizio Ticketing"),
    ("3.5.1.2","Aggiornamento Documentazione API con endpoint RBS 3.5.1",
     "3 Aggiornamento Sistema Esistente > 3.5 Servizio Ticketing"),
    ("3.5.1.3","Testing della funzionalità RBS 3.5.1",
     "3 Aggiornamento Sistema Esistente > 3.5 Servizio Ticketing"),
    ("3.5.2.1","Aggiornamento endpoint di checkout per supportare RBS 3.5.2 nel servizio Ticketing",
     "3 Aggiornamento Sistema Esistente > 3.5 Servizio Ticketing"),
    ("3.5.2.2","Aggiornamento Documentazione API endpoint di checkout",
     "3 Aggiornamento Sistema Esistente > 3.5 Servizio Ticketing"),
    ("3.5.2.3","Testing della funzionalità di checkout aggiornata con RBS 3.5.2",
     "3 Aggiornamento Sistema Esistente > 3.5 Servizio Ticketing"),
    ("3.5.3.1","Aggiunta endpoint RBS 3.5.3 al servizio Ticketing",
     "3 Aggiornamento Sistema Esistente > 3.5 Servizio Ticketing"),
    ("3.5.3.2","Aggiornamento Documentazione API con endpoint RBS 3.5.3",
     "3 Aggiornamento Sistema Esistente > 3.5 Servizio Ticketing"),
    ("3.5.3.3","Testing della funzionalità RBS 3.5.3",
     "3 Aggiornamento Sistema Esistente > 3.5 Servizio Ticketing"),
    # 3.6 Frontend Web
    ("3.6.1","Aggiornamento schermata di registrazione/login con supporto OAuth2",
     "3 Aggiornamento Sistema Esistente > 3.6 Frontend Web"),
    ("3.6.2","Integrazione API servizio di geolocalizzazione",
     "3 Aggiornamento Sistema Esistente > 3.6 Frontend Web"),
    ("3.6.3","Supporto nuove notifiche",
     "3 Aggiornamento Sistema Esistente > 3.6 Frontend Web"),
    ("3.6.4","Supporto inserimento codice sconto nella schermata di checkout",
     "3 Aggiornamento Sistema Esistente > 3.6 Frontend Web"),
    ("3.6.5","Design schermata per i piani di abbonamento",
     "3 Aggiornamento Sistema Esistente > 3.6 Frontend Web"),
    ("3.6.6","Implementazione schermata per i piani di abbonamento",
     "3 Aggiornamento Sistema Esistente > 3.6 Frontend Web"),
    ("3.6.7","Integrazione sistema di abbonamento nelle schermate del sito, con aggiunta schermata ad hoc per le statistiche di performance per le organizzazioni",
     "3 Aggiornamento Sistema Esistente > 3.6 Frontend Web"),
    # 3.7 Requisiti Non Funzionali
    ("3.7.1.1","Verificare che i nuovi endpoint e quelli aggiornati rispettino il requisito di performance",
     "3 Aggiornamento Sistema Esistente > 3.7 Requisiti Non Funzionali"),
    ("3.7.2.1","Verificare che il frontend web sia responsive",
     "3 Aggiornamento Sistema Esistente > 3.7 Requisiti Non Funzionali"),
]

PARTICIPANTS   = ["Alice Alfonsi", "Federico Bravetti", "Tommaso Brini"]
SAVE_FILE    = "stime_delphi.json"
OUTPUT_MD    = "stime_delphi.md"
OUTPUT_DOCX  = "stime_delphi.docx"
OUTPUT_GANTT = "stime_delphi_piano.docx"
SEP          = "─" * 72

DEFAULT_PROJECT_START = datetime(2026, 6, 1, 6, 0, 0)
DATE_INPUT_FORMATS    = ["%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d", "%d/%m/%y", "%d.%m.%Y"]
DATE_OUTPUT_FORMAT    = "%d/%m/%Y"


# ── I/O HELPERS ───────────────────────────────────────────────────────────

def load_state():
    if os.path.exists(SAVE_FILE):
        with open(SAVE_FILE, encoding="utf-8") as f:
            return json.load(f)
    return {}

def save_state(state):
    with open(SAVE_FILE, "w", encoding="utf-8") as f:
        json.dump(state, f, indent=2, ensure_ascii=False)

def ask_float(prompt):
    while True:
        try:
            return float(input(prompt).strip().replace(",", "."))
        except ValueError:
            print("  !! Inserisci un numero valido (es. 0.5 oppure 2).")

def ask_menu(options):
    """options: dict {tasto: etichetta}  (tasto '' = Invio)"""
    labels = "   ".join(
        f"[{'Invio' if k == '' else k}] {v}"
        for k, v in options.items()
    )
    print(f"  {labels}")
    valid = set(options.keys())
    while True:
        raw = input("  → ").strip().lower()
        if raw in valid:
            return raw
        opts = ", ".join(f"[{'Invio' if k == '' else k}]" for k in options)
        print(f"  !! '{raw}' non riconosciuto. Opzioni valide: {opts}")

def show_rounds(rounds):
    if not rounds:
        return
    header = f"  {'Round':>5}  {'Alice':>8}  {'Federico':>10}  {'Tommaso':>9}  {'Media':>7}  {'Mediana':>8}"
    print(header)
    print("  " + "-" * (len(header) - 2))
    for r in rounds:
        print(f"  {r['round']:>5}  {r['alice']:>8.2f}  {r['federico']:>10.2f}  "
              f"{r['tommaso']:>9.2f}  {r['mean']:>7.2f}  {r['median']:>8.2f}")
    print("  (valori in ore/u)")


# ── CORE LOOP ─────────────────────────────────────────────────────────────

def collect_round(rounds, round_num):
    """Chiede le stime ai 3 partecipanti e restituisce il record del round."""
    print(f"\n  ══ Round {round_num} ══")
    vals = []
    for p in PARTICIPANTS:
        v = ask_float(f"    {p:<22} (ore/u): ")
        vals.append(v)

    mean_val   = statistics.mean(vals)
    median_val = statistics.median(vals)
    record = {
        "round":    round_num,
        "alice":    vals[0],
        "federico": vals[1],
        "tommaso":  vals[2],
        "mean":     round(mean_val,   4),
        "median":   round(median_val, 4),
    }
    print(f"\n  → Media: {mean_val:.2f} ore/u   |   Mediana: {median_val:.2f} ore/u")
    return record


def process_activity(act_id, act_name, section, rounds):
    """
    Gestisce tutti i round di una singola attività.
    Restituisce True per continuare, False per uscire dallo script.
    """
    total = len(ACTIVITIES)
    idx   = next(i for i, a in enumerate(ACTIVITIES) if a[0] == act_id) + 1

    print(f"\n{SEP}")
    print(f"  [{idx}/{total}]  {act_id}  –  {act_name}")
    print(f"  {section}")
    print(SEP)

    if rounds:
        print(f"\n  Round già registrati:")
        show_rounds(rounds)

        choice = ask_menu({"": "nuovo round", "s": "salta attività", "q": "esci e salva"})
        if choice == "q":
            return False
        if choice == "s":
            return True

    while True:
        round_num = len(rounds) + 1
        record    = collect_round(rounds, round_num)
        rounds.append(record)

        print()
        show_rounds(rounds)

        choice = ask_menu({"": "altro round", "n": "prossima attività", "q": "esci e salva"})
        if choice == "q":
            return False
        if choice == "n":
            return True


# ── SCHEDULE HELPERS ──────────────────────────────────────────────────────

def parse_user_date(s):
    s = (s or "").strip()
    if not s:
        return None
    for fmt in DATE_INPUT_FORMATS:
        try:
            return datetime.strptime(s, fmt)
        except ValueError:
            continue
    return None

def format_user_date(dt):
    return dt.strftime(DATE_OUTPUT_FORMAT) if dt else ""

def _count_work_days(start_dt, end_dt):
    """Conta i giorni lavorativi (lun-ven) da start_dt a end_dt, inclusivi."""
    days = 0
    current = start_dt.date() if hasattr(start_dt, "date") else start_dt
    end = end_dt.date() if hasattr(end_dt, "date") else end_dt
    while current <= end:
        if current.weekday() < 5:
            days += 1
        current += timedelta(days=1)
    return max(1, days)


def resolve_schedule(durata_raw, inizio_raw, fine_raw,
                     default_duration, default_start):
    """Risolve (durata_giorni, inizio, fine) con la formula come unica fonte
    di verità per la DURATA. Le date sono input dell'utente e non vengono
    MAI ricalcolate: vengono ritornate esattamente come lette (None se vuote).
    """
    duration = max(1, int(default_duration))
    inizio   = parse_user_date(inizio_raw)
    fine     = parse_user_date(fine_raw)
    return duration, inizio, fine


# ── OUTPUT ────────────────────────────────────────────────────────────────

def generate_markdown(state):
    ts = datetime.now().strftime("%Y-%m-%d %H:%M")
    lines = [
        "# Stime Delphi – EvenToNight",
        "",
        f"_Generato il {ts}_  ",
        "_Unità: ore/uomo (ore/u)_",
        "",
    ]

    current_top = None
    for act_id, act_name, section in ACTIVITIES:
        if act_id not in state or not state[act_id]:
            continue

        top = section.split(" > ")[0]
        if top != current_top:
            current_top = top
            lines += [f"## {top}", ""]

        rounds = state[act_id]
        lines += [
            f"### Attività {act_id} – {act_name}",
            "",
            "| Round | Alice Alfonsi | Federico Bravetti | Tommaso Brini | Media | Mediana |",
            "|------:|:-------------:|:-----------------:|:-------------:|------:|--------:|",
        ]
        for r in rounds:
            lines.append(
                f"| {r['round']} "
                f"| {r['alice']:.1f} "
                f"| {r['federico']:.1f} "
                f"| {r['tommaso']:.1f} "
                f"| {r['mean']:.1f} "
                f"| {r['median']:.1f} |"
            )
        lines.append("")

    return "\n".join(lines)


# ── WORD ──────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def _bold_cell(cell, text, font_size=10, color_hex=None, bg_hex=None):
    cell.text = ""
    p   = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    if bg_hex:
        _set_cell_bg(cell, bg_hex)

def _plain_cell(cell, text, font_size=10):
    cell.text = ""
    p   = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string("595959")

def _plain_cell_left(cell, text, font_size=10):
    cell.text = ""
    p   = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string("595959")

def _load_existing_gantt_fields(filepath):
    """Legge dal piano già esistente i campi PREDECESSORE/ASSEGNATO/DATA INIZIO/DATA FINE.
    Restituisce dict {nome_task: [pred, assegnato, d_inizio, d_fine]}."""
    if not os.path.exists(filepath):
        return {}
    try:
        doc = Document(filepath)
        if not doc.tables:
            return {}
        preserved = {}
        for row in doc.tables[0].rows[1:]:   # salta header
            cells = row.cells
            if len(cells) < 9:
                continue
            task_name = cells[2].text.strip()
            if not task_name:
                continue
            preserved[task_name] = [
                cells[4].text.strip(),   # DURATA
                cells[5].text.strip(),   # PREDECESSORE
                cells[6].text.strip(),   # ASSEGNATO
                cells[7].text.strip(),   # DATA INIZIO
                cells[8].text.strip(),   # DATA FINE
            ]
        return preserved
    except Exception:
        return {}

def _set_cell_valign_center(cell):
    tcPr   = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)

def generate_word(state):
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    title = doc.add_heading("Stime Delphi – EvenToNight", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor.from_string("1C4587")

    sub = doc.add_paragraph(
        f"Generato il {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  "
        f"Unità: ore/uomo (ore/u)"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor.from_string("666666")

    HEADER_BG = "D6E4F0"   # azzurro chiaro
    HEADER_FG = "1F4E79"   # blu navy
    ALT_BG    = "EBF3FB"   # azzurro molto chiaro (righe alternate)
    COL_W     = [Cm(1.8), Cm(3.5), Cm(3.5), Cm(3.5), Cm(2.5), Cm(2.7)]
    COL_HEADS = [
        "Round",
        "Alice Alfonsi\n(ore/u)",
        "Federico Bravetti\n(ore/u)",
        "Tommaso Brini\n(ore/u)",
        "Media\n(ore/u)",
        "Mediana\n(ore/u)",
    ]

    current_top = None
    for act_id, act_name, section in ACTIVITIES:
        if act_id not in state or not state[act_id]:
            continue

        top = section.split(" > ")[0]
        if top != current_top:
            current_top = top
            h1 = doc.add_heading(top, level=1)
            h1.runs[0].font.color.rgb = RGBColor.from_string("1155CC")

        h2 = doc.add_heading(f"Attività {act_id} – {act_name}", level=2)
        h2.runs[0].font.color.rgb = RGBColor.from_string("434343")

        rounds = state[act_id]
        tbl = doc.add_table(rows=1 + len(rounds), cols=6)
        tbl.style = "Table Grid"

        for row in tbl.rows:
            for c_idx, cell in enumerate(row.cells):
                cell.width = COL_W[c_idx]

        for c_idx, label in enumerate(COL_HEADS):
            _bold_cell(tbl.rows[0].cells[c_idx], label,
                       font_size=9, color_hex=HEADER_FG, bg_hex=HEADER_BG)

        for r_idx, r in enumerate(rounds):
            row = tbl.rows[r_idx + 1]
            bg  = ALT_BG if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, v in enumerate([
                str(r["round"]),
                f"{r['alice']:.1f}",
                f"{r['federico']:.1f}",
                f"{r['tommaso']:.1f}",
                f"{r['mean']:.1f}",
                f"{r['median']:.1f}",
            ]):
                _plain_cell(row.cells[c_idx], v, font_size=9)
                _set_cell_bg(row.cells[c_idx], bg)

        doc.add_paragraph()

    doc.save(OUTPUT_DOCX)


# ── PIANO ATTIVITÀ (tabella gantt-style) ──────────────────────────────────

def generate_gantt_word(state, ore_giorno, efficienza):
    # preserva i campi manuali dal file esistente (se presente)
    preserved = _load_existing_gantt_fields(OUTPUT_GANTT)

    # raccoglie solo attività con stime, in ordine WBS
    collected = []
    for act_id, act_name, section in ACTIVITIES:
        if act_id not in state or not state[act_id]:
            continue
        top    = section.split(" > ")[0]
        median = state[act_id][-1]["median"]   # ultimo round
        collected.append((act_id, act_name, top, median))

    if not collected:
        print("  Nessuna stima disponibile.")
        return

    # gruppi per sezione macro (per merge celle SUBSYSTEM)
    groups = []
    cur_top, cur_start = None, 0
    for i, (_, _, top, _) in enumerate(collected):
        if top != cur_top:
            if cur_top is not None:
                groups.append((cur_top, cur_start, i - 1))
            cur_top, cur_start = top, i
    groups.append((cur_top, cur_start, len(collected) - 1))

    doc = Document()

    # ── landscape A4 ──
    sec = doc.sections[0]
    sec.orientation   = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin    = sec.bottom_margin = Cm(1.5)
    sec.left_margin   = sec.right_margin  = Cm(1.5)

    # ── titolo ──
    t = doc.add_heading("Piano delle Attività – EvenToNight", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.color.rgb = RGBColor.from_string("1C4587")
    eff_pct = efficienza * 100
    eff_str = f"{int(eff_pct)}%" if eff_pct == int(eff_pct) else f"{eff_pct:g}%"
    ore_str = f"{int(ore_giorno)}" if ore_giorno == int(ore_giorno) else f"{ore_giorno:g}"
    sub = doc.add_paragraph(
        f"Generato il {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  "
        f"Durata = ⌈Effort / ({ore_str} ore × {eff_str})⌉ giorni"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor.from_string("666666")

    HEADER_BG = "D6E4F0"
    HEADER_FG = "1F4E79"
    SUBSYS_BG = "BDD7EE"
    ALT_BG    = "EBF3FB"

    # larghezze colonne (landscape ~14500 twips disponibili)
    COL_W = [Cm(1.0), Cm(2.8), Cm(7.5), Cm(1.6), Cm(1.6),
             Cm(2.5), Cm(2.5), Cm(2.0), Cm(2.0)]
    COL_H = [
        "TASK ID", "SUBSYSTEM", "TASK",
        "EFFORT\n(ore/u)", "DURATA\n(gg)",
        "PREDECESSORE", "ASSEGNATO", "DATA INIZIO", "DATA FINE",
    ]

    n   = len(collected)
    tbl = doc.add_table(rows=1 + n, cols=9)
    tbl.style = "Table Grid"

    for row in tbl.rows:
        for c, cell in enumerate(row.cells):
            cell.width = COL_W[c]

    # ── header ──
    for c, label in enumerate(COL_H):
        _bold_cell(tbl.rows[0].cells[c], label,
                   font_size=8, color_hex=HEADER_FG, bg_hex=HEADER_BG)

    # ── righe dati ──
    for i, (act_id, act_name, top, median) in enumerate(collected):
        row    = tbl.rows[i + 1]
        bg     = ALT_BG if i % 2 == 0 else "FFFFFF"
        effort = median
        default_duration = compute_duration(median, ore_giorno, efficienza)

        prev = preserved.get(act_name, ["", "", "", "", ""])
        resolved_dur, resolved_inizio, resolved_fine = resolve_schedule(
            prev[0], prev[3], prev[4],
            default_duration=default_duration,
            default_start=DEFAULT_PROJECT_START,
        )
        durata_val = str(resolved_dur)
        inizio_str = format_user_date(resolved_inizio)
        fine_str   = format_user_date(resolved_fine)

        for c, (v, style) in enumerate([
            (str(i + 1),        "center"),
            (top,               "subsys"),   # placeholder, verrà sostituito dal merge
            (act_name,          "left"),
            (f"{effort:.1f}",   "center"),
            (durata_val,        "center"),   # DURATA (risolta da date se presenti)
            (prev[1],           "center"),   # PREDECESSORE (preservato)
            (prev[2],           "center"),   # ASSEGNATO (preservato)
            (inizio_str,        "center"),   # DATA INIZIO (auto-completata)
            (fine_str,          "center"),   # DATA FINE (auto-completata)
        ]):
            cell = row.cells[c]
            if style == "subsys":
                _bold_cell(cell, v, font_size=8,
                           color_hex=HEADER_FG, bg_hex=SUBSYS_BG)
            elif style == "left":
                _plain_cell_left(cell, v, font_size=8)
                _set_cell_bg(cell, bg)
            else:
                _plain_cell(cell, v, font_size=8)
                _set_cell_bg(cell, bg)

    # ── merge celle SUBSYSTEM per gruppo ──
    for top_name, start, end in groups:
        if end > start:
            tbl.cell(start + 1, 1).merge(tbl.cell(end + 1, 1))
        mc       = tbl.cell(start + 1, 1)
        mc.text  = ""
        p        = mc.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run      = p.add_run(top_name)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(HEADER_FG)
        _set_cell_bg(mc, SUBSYS_BG)
        _set_cell_valign_center(mc)

    doc.save(OUTPUT_GANTT)


# ── GANTT JSON ────────────────────────────────────────────────────────────

OUTPUT_GANTT_JSON         = "eventonight.gantt"
OUTPUT_GANTT_JSON_NOPRED  = "eventonight_nopred.gantt"
OUTPUT_GANTT_JSON_COLORED = "eventonight_colored.gantt"

_ADVANCED_DEFAULT = {
    "columns": [
        {"name": "Task ID",    "width": "70",  "show": True},
        {"name": "Task Name",  "width": "350", "show": True},
        {"name": "Start Date", "width": "130", "show": False},
        {"name": "End Date",   "width": "130", "show": False},
        {"name": "Duration",   "width": "130", "show": False},
        {"name": "Progress %", "width": "150", "show": False},
        {"name": "Dependency", "width": "150", "show": False},
        {"name": "Resources",  "width": "200", "show": False},
        {"name": "Color",      "width": "100", "show": False},
    ],
    "zoomLevel": 0,
    "timezone": "Europe/Rome",
    "timezoneOffset": -120,
    "dependencyConflict": "Add Offset to Dependency",
    "dateFormat": "yyyy-MM-dd",
    "timeFormat": "HH:mm",
    "firstDayOfWeek": 0,
    "workWeek": ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"],
    "workTime": [{"from": 8, "to": 12}, {"from": 13, "to": 17}],
    "holidays": [],
}

_RESOURCE_NAMES = {1: "Alice Alfonsi", 2: "Federico Bravetti", 3: "Tommaso Brini"}
# lookup case-insensitive solo sul nome completo (nessuna sigla/forma breve)
_RESOURCE_LOOKUP = {name.lower(): rid for rid, name in _RESOURCE_NAMES.items()}
# top-level resources: formato esatto del tool (resourceId = nome stringa)
_GANTT_RESOURCES = [
    {"resourceId": name, "resourceName": name}
    for name in _RESOURCE_NAMES.values()
]

def _split_assignees(assegnato_raw):
    """Divide il campo ASSEGNATO in nomi (separatori , ; /) e restituisce
    la lista degli ID risorsa corrispondenti, nell'ordine di apparizione."""
    seen, ids = set(), []
    for chunk in (assegnato_raw or "").replace(";", ",").replace("/", ",").split(","):
        rid = _RESOURCE_LOOKUP.get(chunk.strip().lower())
        if rid and rid not in seen:
            seen.add(rid)
            ids.append(rid)
    return ids

def _parse_resources(assegnato_raw):
    """Converte 'Alice Alfonsi, Federico Bravetti' → lista resource refs."""
    return [
        {"resourceId": _RESOURCE_NAMES[rid], "resourceName": _RESOURCE_NAMES[rid], "unit": 100}
        for rid in _split_assignees(assegnato_raw)
    ]

# colori numerici proprietari del tool (solo codici della palette nativa).
# Per le combinazioni si usa il colore del primo assegnatario in ordine
# alfabetico/ID, perché i codici 4xx-7xx non corrispondono a colori validi.
_ASSIGNMENT_COLORS = {
    frozenset([1]):       "121",   # Alice solo
    frozenset([2]):       "211",   # Federico solo
    frozenset([3]):       "301",   # Tommaso solo
    frozenset([1, 2]):    "121",   # Alice + Federico  → colore Alice
    frozenset([1, 3]):    "121",   # Alice + Tommaso   → colore Alice
    frozenset([2, 3]):    "211",   # Federico + Tommaso → colore Federico
    frozenset([1, 2, 3]): "121",   # tutti              → colore Alice
}

def _assignment_color(assegnato_raw):
    ids = set(_split_assignees(assegnato_raw))
    return _ASSIGNMENT_COLORS.get(frozenset(ids), "") if ids else ""

def compute_duration(effort, ore_giorno, efficienza):
    """Durata = ⌈ effort / (ore_giorno * efficienza) ⌉.

    - effort: ore/uomo (mediana dei round Delphi)
    - ore_giorno: ore lavorative al giorno
    - efficienza: frazione (es. 0.8 = 80%)
    """
    divisor = ore_giorno * efficienza
    if divisor <= 0:
        return 1
    return max(1, math.ceil(effort / divisor))

def _clean_predecessor(raw):
    """Normalizza predecessori al formato 'NNN SS' (Start-to-Start).

    Estrae i numeri da qualsiasi formato (plain, 2FS, 3SS…) e ri-appone SS
    così le frecce dipendenza sono di tipo Start-to-Start.
    """
    import re
    parts = re.split(r"[,;\s]+", (raw or "").strip())
    result = []
    for p in parts:
        m = re.match(r"^(\d+)", p.strip())
        if m:
            result.append(m.group(1) + "SS")
    return ",".join(result)

def generate_gantt_json(state, ore_giorno, efficienza):
    # preserva blocco advanced dall'esistente (zoom, timezone, workTime…)
    advanced = _ADVANCED_DEFAULT
    if os.path.exists(OUTPUT_GANTT_JSON):
        try:
            with open(OUTPUT_GANTT_JSON, encoding="utf-8") as f:
                existing = json.load(f)
            if "advanced" in existing:
                advanced = existing["advanced"]
        except Exception:
            pass

    # legge tutti i campi dal piano Word (a questo punto già risolto)
    word_fields  = _load_existing_gantt_fields(OUTPUT_GANTT)
    predecessors = {n: _clean_predecessor(f[1]) for n, f in word_fields.items() if f[1].strip()}

    # raccoglie attività stimate in ordine WBS, assegna ID sequenziali 1,2,3…
    # i subtask usano questi ID; le sezioni macro ricevono ID 1001,1002,1003
    # così il campo Predecessor nel gantt combacia con TASK ID del Word
    collected   = []   # [(act_id, act_name, top, duration, inizio, fine)]
    groups      = {}   # top -> [indice in collected]
    group_order = []
    for act_id, act_name, section in ACTIVITIES:
        if act_id not in state or not state[act_id]:
            continue
        top              = section.split(" > ")[0]
        median           = state[act_id][-1]["median"]
        default_duration = compute_duration(median, ore_giorno, efficienza)
        prev             = word_fields.get(act_name, ["", "", "", "", ""])
        duration, inizio, fine = resolve_schedule(
            prev[0], prev[3], prev[4],
            default_duration=default_duration,
            default_start=DEFAULT_PROJECT_START,
        )
        if top not in groups:
            groups[top] = []
            group_order.append(top)
        groups[top].append(len(collected))
        collected.append((act_id, act_name, top, duration, inizio, fine))

    if not collected:
        print("  Nessuna stima disponibile per il .gantt.")
        return

    def iso(dt):
        return dt.strftime("%Y-%m-%dT%H:%M:%S.000Z")

    n_subtasks = len(collected)          # 1-based IDs 1..n_subtasks per le foglie
    data = []
    for macro_idx, top in enumerate(group_order):
        top_id   = n_subtasks + 1 + macro_idx   # IDs immediatamente dopo i subtask, nessun gap
        subtasks = []
        starts   = []
        ends     = []

        for seq_idx in groups[top]:
            act_id, act_name, _, duration, inizio, fine = collected[seq_idx]
            sub_task_id = seq_idx + 1              # ID 1-based, coincide col TASK ID Word
            pred_str    = predecessors.get(act_name, "")
            assegnato   = word_fields.get(act_name, ["", "", "", "", ""])[2]

            # Fallback SOLO per il gantt JSON quando il docx non ha date:
            # usa default_start e propaga la durata. Non scrive nulla nel docx.
            inizio_dt = inizio or fine or DEFAULT_PROJECT_START
            fine_dt   = fine or (inizio_dt + timedelta(days=duration - 1))

            start_dt = inizio_dt.replace(hour=6,  minute=0, second=0, microsecond=0)
            end_dt   = fine_dt.replace(  hour=15, minute=0, second=0, microsecond=0)
            starts.append(start_dt)
            ends.append(end_dt)

            # Duration in giorni lavorativi (lun-ven) per allinearsi al workWeek
            # del tool: così il tool calcola End = ultimo giorno lavorativo
            # coincidente con la DATA FINE del piano Word.
            work_duration = _count_work_days(start_dt, end_dt)

            subtasks.append({
                "TaskID":       sub_task_id,
                "TaskName":     f"{act_id} {act_name}",
                "StartDate":    iso(start_dt),
                "EndDate":      iso(end_dt),
                "Duration":     work_duration,
                "Predecessor":  pred_str,
                "resources":    _parse_resources(assegnato),
                "Progress":     0,
                "color":        "",
                "info":         "<p><br></p>",
                "DurationUnit": "day",
                "isManual":     True,
                "_color":       _assignment_color(assegnato),   # usato dal colored
            })

        macro_start = min(starts)
        macro_end   = max(ends)
        macro_dur   = _count_work_days(macro_start, macro_end)

        data.append({
            "TaskID":       top_id,
            "TaskName":     top,
            "StartDate":    iso(macro_start),
            "EndDate":      iso(macro_end),
            "Duration":     macro_dur,
            "Predecessor":  None,
            "resources":    [],
            "Progress":     0,
            "color":        "",
            "info":         "<p><br></p>",
            "DurationUnit": "day",
            "subtasks":     subtasks,
        })

    output = {
        "data":             data,
        "resources":        _GANTT_RESOURCES,
        "projectStartDate": None,
        "projectEndDate":   None,
        "advanced":         advanced,
    }

    # rimuove il campo temporaneo _color prima di scrivere il gantt principale
    import copy
    def _strip_tmp(data):
        d = copy.deepcopy(data)
        for group in d["data"]:
            for sub in group.get("subtasks", []):
                sub.pop("_color", None)
        return d

    with open(OUTPUT_GANTT_JSON, "w", encoding="utf-8") as f:
        json.dump(_strip_tmp(output), f, ensure_ascii=False, separators=(",", ":"))

    # variante senza predecessori
    output_nopred = _strip_tmp(output)
    for group in output_nopred["data"]:
        group["Predecessor"] = None
        for sub in group.get("subtasks", []):
            sub["Predecessor"] = ""
    with open(OUTPUT_GANTT_JSON_NOPRED, "w", encoding="utf-8") as f:
        json.dump(output_nopred, f, ensure_ascii=False, separators=(",", ":"))

    # variante colorata per assegnatario, senza predecessori
    output_colored = copy.deepcopy(output)
    for group in output_colored["data"]:
        group["Predecessor"] = None
        for sub in group.get("subtasks", []):
            sub["color"] = sub.pop("_color", "")
            sub["Predecessor"] = ""
    with open(OUTPUT_GANTT_JSON_COLORED, "w", encoding="utf-8") as f:
        json.dump(output_colored, f, ensure_ascii=False, separators=(",", ":"))


# ── MAIN ──────────────────────────────────────────────────────────────────

def _ask_duration_params():
    print(f"\n{SEP}")
    print("  Generazione file di output...")
    print(f"  Il piano attività calcola DURATA = ⌈Effort / (ore × efficienza)⌉ giorni.")
    ore_raw = input("  Ore lavorative al giorno [Invio = 8]: ").strip().replace(",", ".")
    if ore_raw:
        try:
            ore_giorno = float(ore_raw)
            if ore_giorno <= 0:
                raise ValueError
        except ValueError:
            print("  Valore non valido, uso 8.")
            ore_giorno = 8.0
    else:
        ore_giorno = 8.0

    eff_raw = input("  Efficienza % [Invio = 100]: ").strip().replace(",", ".").rstrip("%")
    if eff_raw:
        try:
            eff_val = float(eff_raw)
            if eff_val <= 0:
                raise ValueError
            efficienza = eff_val / 100 if eff_val > 1 else eff_val
        except ValueError:
            print("  Valore non valido, uso 100%.")
            efficienza = 1.0
    else:
        efficienza = 1.0

    ore_str = f"{int(ore_giorno)}" if ore_giorno == int(ore_giorno) else f"{ore_giorno:g}"
    eff_pct = efficienza * 100
    eff_str = f"{int(eff_pct)}%" if eff_pct == int(eff_pct) else f"{eff_pct:g}%"
    print(f"  Confermato: ore/giorno = {ore_str}, efficienza = {eff_str}")
    return ore_giorno, efficienza

def _generate_all(state, ore_giorno, efficienza):
    md = generate_markdown(state)
    with open(OUTPUT_MD, "w", encoding="utf-8") as f:
        f.write(md)
    generate_word(state)
    generate_gantt_word(state, ore_giorno, efficienza)
    generate_gantt_json(state, ore_giorno, efficienza)

    done_now = sum(1 for a in ACTIVITIES if a[0] in state and state[a[0]])
    print(f"\n{SEP}")
    print(f"  Salvataggio completato.")
    print(f"  Attività con stime: {done_now}/{len(ACTIVITIES)}")
    print(f"  JSON         → {SAVE_FILE}")
    print(f"  MD           → {OUTPUT_MD}")
    print(f"  Stime Word   → {OUTPUT_DOCX}")
    print(f"  Piano Word   → {OUTPUT_GANTT}")
    print(f"  Gantt        → {OUTPUT_GANTT_JSON}")
    print(f"  Gantt nopred → {OUTPUT_GANTT_JSON_NOPRED}")
    print(f"  Gantt colored→ {OUTPUT_GANTT_JSON_COLORED}")
    print(SEP)

def main():
    state = load_state()

    done_count = sum(1 for a in ACTIVITIES if a[0] in state and state[a[0]])
    total      = len(ACTIVITIES)

    print(f"""
{SEP}
  STIME DELPHI – Espansione EvenToNight
  {done_count}/{total} attività già con dati  |  Unità: ore/uomo (ore/u)
  Comandi: [Invio] = conferma/altro round   [n] = prossima   [q] = esci e salva
{SEP}""")

    if done_count > 0:
        print()
        scelta = ask_menu({
            "": "inserisci/modifica stime",
            "r": "rigenera output (Word + Gantt) senza modificare stime",
        })
        if scelta == "r":
            ore_giorno, efficienza = _ask_duration_params()
            _generate_all(state, ore_giorno, efficienza)
            return

    for act_id, act_name, section in ACTIVITIES:
        rounds = state.setdefault(act_id, [])
        keep_going = process_activity(act_id, act_name, section, rounds)
        save_state(state)
        if not keep_going:
            break

    ore_giorno, efficienza = _ask_duration_params()
    _generate_all(state, ore_giorno, efficienza)


if __name__ == "__main__":
    main()
